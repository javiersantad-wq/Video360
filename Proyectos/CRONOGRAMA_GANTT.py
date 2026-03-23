"""
Generador de Cronograma Maestro Tipo Gantt
Similar al documento de referencia
"""

import geopandas as gpd
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime, timedelta

BASE_DIR = r"C:\Users\ed\.openclaw\workspace\Proyectos"

def load_blocks(path):
    """Carga shapefile y resume por bloque"""
    gdf = gpd.read_file(path)
    results = []
    for bloque in sorted(gdf['bloque'].unique()):
        block_data = gdf[gdf['bloque'] == bloque]
        results.append({
            'bloque': int(bloque),
            'segmentos': len(block_data),
            'km': block_data['km'].sum(),
        })
    return pd.DataFrame(results)

def add_color(cell, hex_color):
    """Agrega color a celda"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_width(cell, width_cm):
    """Ajusta ancho de celda"""
    cell.width = width_cm

def create_gantt_document(df_ts, df_vt, start_date, output_path):
    """Crea documento Gantt combinado para ambos proyectos"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Cronograma Maestro: Proyectos de Vuelo LiDAR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('PROYECTO TAMPICO - SAN LUIS y VERACRUZ - TAMPICO')
    doc.add_paragraph()
    
    # Calculate all dates
    # Project 1: 9 bloques = 18 days (vuelo+proceso) + 2 setup = 20 days
    # Project 2: 11 bloques = 22 days (vuelo+proceso) + 1 setup = 23 days
    # Total: 43 days + buffer
    
    # Generate date range starting March 30
    current_date = datetime(2026, 3, 30)
    all_dates = []
    
    # Project 1
    dates_ts = []
    current_date_ts = current_date
    for _, row in df_ts.iterrows():
        # Vuelo day
        dates_ts.append((current_date_ts, f"Bloque {row['bloque']}", "Vuelo", row['km']))
        current_date_ts += timedelta(days=1)
        # Proceso day
        dates_ts.append((current_date_ts, f"Bloque {row['bloque']}", "Proceso", 0))
        current_date_ts += timedelta(days=1)
    
    end_date_1 = current_date_ts
    all_dates.extend(dates_ts)
    
    # Gap day
    all_dates.append((current_date_ts, "Transicion", "Setup", 0))
    current_date_ts += timedelta(days=1)
    
    # Project 2
    dates_vt = []
    for _, row in df_vt.iterrows():
        # Vuelo day
        dates_vt.append((current_date_ts, f"Bloque {row['bloque']}", "Vuelo", row['km']))
        current_date_ts += timedelta(days=1)
        # Proceso day
        dates_vt.append((current_date_ts, f"Bloque {row['bloque']}", "Proceso", 0))
        current_date_ts += timedelta(days=1)
    
    end_date_2 = current_date_ts
    all_dates.extend(dates_vt)
    
    # Final processing days
    for i in range(5):
        all_dates.append((current_date_ts, f"Procesamiento final {i+1}", "Proceso", 0))
        current_date_ts += timedelta(days=1)
    
    final_date = current_date_ts
    
    # === PROJECT 1 TABLE ===
    doc.add_heading('PROYECTO 1: TAMPICO - SAN LUIS', level=1)
    doc.add_paragraph(f'Fecha de inicio: {datetime(2026, 3, 30).strftime("%d de %B de %Y")}')
    doc.add_paragraph(f'Bloques: {len(df_ts)} | Kilometros: {df_ts["km"].sum():.1f} km')
    doc.add_paragraph()
    
    # Create table
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    
    hdr1 = table1.rows[0].cells
    hdr1[0].text = 'Dia'
    hdr1[1].text = 'Fecha'
    hdr1[2].text = 'Actividad'
    hdr1[3].text = 'km'
    
    for para in hdr1[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
    for para in hdr1[1].paragraphs:
        for run in para.runs:
            run.font.bold = True
    for para in hdr1[2].paragraphs:
        for run in para.runs:
            run.font.bold = True
    for para in hdr1[3].paragraphs:
        for run in para.runs:
            run.font.bold = True
    
    day_num = 1
    for date, activity, tipo, km in dates_ts:
        row = table1.add_row()
        row.cells[0].text = str(day_num)
        row.cells[1].text = date.strftime('%d/%m')
        row.cells[2].text = activity
        row.cells[3].text = f'{km:.0f}' if km > 0 else ''
        
        if tipo == "Vuelo":
            add_color(row.cells[2], 'D9E2F3')
        else:
            add_color(row.cells[2], 'E2EFDA')
        
        day_num += 1
    
    doc.add_paragraph()
    
    # === PROJECT 2 TABLE ===
    doc.add_heading('PROYECTO 2: VERACRUZ - TAMPICO', level=1)
    doc.add_paragraph(f'Fecha de inicio: {end_date_1.strftime("%d de %B de %Y")}')
    doc.add_paragraph(f'Bloques: {len(df_vt)} | Kilometros: {df_vt["km"].sum():.1f} km')
    doc.add_paragraph()
    
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Dia'
    hdr2[1].text = 'Fecha'
    hdr2[2].text = 'Actividad'
    hdr2[3].text = 'km'
    
    for para in hdr2[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
    for para in hdr2[1].paragraphs:
        for run in para.runs:
            run.font.bold = True
    for para in hdr2[2].paragraphs:
        for run in para.runs:
            run.font.bold = True
    for para in hdr2[3].paragraphs:
        for run in para.runs:
            run.font.bold = True
    
    day_num = 1
    # Skip first entry (transition day already counted)
    dates_vt_with_transition = [(end_date_1, "Transicion", "Setup", 0)] + dates_vt
    
    for date, activity, tipo, km in dates_vt_with_transition:
        row = table2.add_row()
        row.cells[0].text = str(day_num)
        row.cells[1].text = date.strftime('%d/%m')
        row.cells[2].text = activity
        row.cells[3].text = f'{km:.0f}' if km > 0 else ''
        
        if tipo == "Vuelo":
            add_color(row.cells[2], 'D9E2F3')
        elif tipo == "Proceso":
            add_color(row.cells[2], 'E2EFDA')
        else:
            add_color(row.cells[2], 'FFF2CC')  # Yellow for setup
        
        day_num += 1
    
    # Final processing
    doc.add_paragraph()
    doc.add_heading('PROCESAMIENTO FINAL', level=1)
    doc.add_paragraph(f'Fecha de inicio: {end_date_2.strftime("%d de %B de %Y")}')
    doc.add_paragraph('5 dias para procesamiento, validacion y entrega de resultados')
    
    # Summary table
    doc.add_paragraph()
    doc.add_heading('RESUMEN GENERAL', level=1)
    
    total_km = df_ts['km'].sum() + df_vt['km'].sum()
    total_bloques = len(df_ts) + len(df_vt)
    total_days = (final_date - datetime(2026, 3, 30)).days
    
    summary = [
        ('Total bloques', str(total_bloques)),
        ('Total kilometros', f'{total_km:.1f} km'),
        ('Dias de vuelo', str(total_bloques)),
        ('Dias de procesamiento', str(total_bloques + 5)),
        ('Dias totales estimados', f'{total_days} dias'),
        ('Fecha de inicio', '30 de Marzo de 2026'),
        ('Fecha de terminacion', final_date.strftime('%d de %B de %Y')),
    ]
    
    for label, value in summary:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(value)
    
    doc.save(output_path)
    print(f'Guardado: {output_path}')
    return final_date

# ============== MAIN ==============

print("Generando cronograma maestro...")

# Load data
ts_path = os.path.join(BASE_DIR, "Tampico_SanLuis", "datos", "Tampico_SanLuis-50km_blocks.shp")
vt_path = os.path.join(BASE_DIR, "Veracruz_Tampico", "datos", "Veracruz-Tampico-50km_blocks_v4.shp")

df_ts = load_blocks(ts_path)
df_vt = load_blocks(vt_path)

print(f"Tampico-San Luis: {len(df_ts)} bloques, {df_ts['km'].sum():.1f} km")
print(f"Veracruz-Tampico: {len(df_vt)} bloques, {df_vt['km'].sum():.1f} km")

output = os.path.join(BASE_DIR, "CRONOGRAMA_MAESTRO_GANTT.docx")
final_date = create_gantt_document(df_ts, df_vt, datetime(2026, 3, 30), output)

print()
print(f"Duracion total: {(final_date - datetime(2026, 3, 30)).days} dias")
print(f"Fecha final: {final_date.strftime('%d de %B de %Y')}")
