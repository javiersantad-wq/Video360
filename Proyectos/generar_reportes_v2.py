"""
Generador de Reportes Word v2 - con datos reales de OSM
"""

import geopandas as gpd
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

BASE_DIR = r"C:\Users\ed\.openclaw\workspace\Proyectos"

def crear_reporte(nombre_proyecto, corredor_path, ejidos_csv, localidades_csv, output_path, descripcion, estados_info):
    """Crea reporte Word con datos reales"""
    
    doc = Document()
    
    # Título
    title = doc.add_heading(f'Reporte de Análisis: {nombre_proyecto}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d de %B de %Y")}')
    doc.add_paragraph()
    
    # Descripción
    doc.add_heading('1. Descripción del Proyecto', level=1)
    doc.add_paragraph(descripcion)
    
    # Estados
    doc.add_heading('2. Estados y Municipios', level=1)
    for estado, info in estados_info.items():
        doc.add_heading(f'  {estado} (CVE: {info["cve"]})', level=2)
        doc.add_paragraph(f'Municipios: {info["municipios"]}')
    
    # Bloques
    doc.add_heading('3. Bloques del Corredor', level=1)
    try:
        corredor = gpd.read_file(corredor_path)
        total_km = corredor['km'].sum() if 'km' in corredor.columns else 0
        doc.add_paragraph(f'Total de bloques: {len(corredor)}')
        doc.add_paragraph(f'Longitud total: {total_km:.1f} km')
        
        # Tabla de bloques
        doc.add_heading('  Lista de Bloques', level=3)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'Nombre', 'km', 'Bloque'
        
        for _, row in corredor.head(20).iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['Name'])[:45]
            cells[1].text = f"{row['km']:.1f}"
            cells[2].text = str(row['bloque'])
        
        if len(corredor) > 20:
            doc.add_paragraph(f'... y {len(corredor) - 20} bloques más')
    except Exception as e:
        doc.add_paragraph(f'Error leyendo bloques: {e}')
    
    # Ejidos
    doc.add_heading('4. Análisis de Predios Ejidales', level=1)
    try:
        ejidos = pd.read_csv(ejidos_csv)
        total_ejidos = len(ejidos)
        
        # Calcular críticos
        criticos = len(ejidos[ejidos['distancia_km'] < 0.5]) if 'distancia_km' in ejidos.columns else 0
        cercanos = len(ejidos[ejidos['distancia_km'] < 1.0]) if 'distancia_km' in ejidos.columns else 0
        
        doc.add_paragraph(f'Total de ejidos a 2 km del corredor: {total_ejidos}')
        doc.add_paragraph(f'Ejidos críticos (< 500m): {criticos}')
        doc.add_paragraph(f'Ejidos muy cercanos (< 1 km): {cercanos}')
        
        # Tabla de ejidos cercanos
        doc.add_heading('  Ejidos Más Cercanos', level=3)
        table_ej = doc.add_table(rows=1, cols=4)
        table_ej.style = 'Table Grid'
        hdr = table_ej.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Nombre', 'Municipio', 'Dist km', 'Nombre Anterior'
        
        # Ordenar por distancia
        if 'distancia_km' in ejidos.columns:
            ejidos_sorted = ejidos.nsmallest(15, 'distancia_km')
        else:
            ejidos_sorted = ejidos.head(15)
        
        for _, row in ejidos_sorted.iterrows():
            cells = table_ej.add_row().cells
            cells[0].text = str(row.get('NOMBRE', 'N/A'))[:35]
            cells[1].text = str(row.get('CVE_MPIO', 'N/A'))
            dist = row.get('distancia_km', 0)
            cells[2].text = f'{dist:.2f}' if pd.notna(dist) else 'N/A'
            cells[3].text = str(row.get('NOM_NA', 'N/A'))[:25]
        
        # Por estado
        if 'CVE_ESTADO' in ejidos.columns:
            doc.add_heading('  Distribución por Estado', level=3)
            for estado, count in ejidos['CVE_ESTADO'].value_counts().items():
                doc.add_paragraph(f'  - Estado {estado}: {count} ejidos')
                
    except Exception as e:
        doc.add_paragraph(f'Error leyendo ejidos: {e}')
    
    # Localidades
    doc.add_heading('5. Localidades Cercanas (OpenStreetMap)', level=1)
    try:
        loc = pd.read_csv(localidades_csv)
        total_loc = len(loc)
        
        # Por tipo
        if 'place' in loc.columns:
            cities = len(loc[loc['place'] == 'city']) if 'city' in loc['place'].values else 0
            towns = len(loc[loc['place'] == 'town']) if 'town' in loc['place'].values else 0
            villages = len(loc[loc['place'] == 'village']) if 'village' in loc['place'].values else 0
            hamlets = len(loc[loc['place'] == 'hamlet']) if 'hamlet' in loc['place'].values else 0
            
            doc.add_paragraph(f'Total de localidades a 5 km: {total_loc}')
            doc.add_paragraph(f'  - Ciudades: {cities}')
            doc.add_paragraph(f'  - Pueblos: {towns}')
            doc.add_paragraph(f'  - Villages: {villages}')
            doc.add_paragraph(f'  - Hamlets: {hamlets}')
        
        # Tabla de principales localidades
        doc.add_heading('  Principales Localidades', level=3)
        table_loc = doc.add_table(rows=1, cols=3)
        table_loc.style = 'Table Grid'
        hdr = table_loc.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'Nombre', 'Tipo', 'Dist km'
        
        if 'dist_to_co' in loc.columns:
            loc_sorted = loc.nsmallest(15, 'dist_to_co')
        else:
            loc_sorted = loc.head(15)
        
        for _, row in loc_sorted.iterrows():
            cells = table_loc.add_row().cells
            cells[0].text = str(row.get('name', 'N/A'))[:35]
            cells[1].text = str(row.get('place', 'N/A'))
            dist = row.get('dist_to_co', 0)
            cells[2].text = f'{dist:.2f}' if pd.notna(dist) else 'N/A'
            
    except FileNotFoundError:
        doc.add_paragraph('Datos de localidades aún no disponibles.')
        doc.add_paragraph('Nota: El servicio de OpenStreetMap está temporalmente sobrecargado.')
        doc.add_paragraph('Se recomienda intentar de nuevo más tarde o descargar manualmente.')
    except Exception as e:
        doc.add_paragraph(f'Error leyendo localidades: {e}')
    
    # Consideraciones
    doc.add_heading('6. Consideraciones para Vuelo de Dron', level=1)
    consideraciones = [
        ('Permisos', 'Consultar Asamblea Ejidal, SEDENA y AFAC'),
        ('Consulta previa', 'Contactar comunidades antes del vuelo'),
        ('Zonas sensibles', 'Identificar zonas urbanas y de alta concentración'),
        ('Seguridad', 'Coordinar con autoridades locales'),
        ('Compensación', 'Algunos ejidos cobran por uso de tierras')
    ]
    for titulo, desc in consideraciones:
        p = doc.add_paragraph()
        p.add_run(f'{titulo}: ').bold = True
        p.add_run(desc)
    
    # Acciones
    doc.add_heading('7. Acciones Recomendadas', level=1)
    acciones = [
        'Revisar archivos CSV para identificar prioridades',
        'Consultar RAN para status legal de ejidos',
        'Contactar autoridades agrarias estatales',
        'Programar reuniones con asambleas ejidales',
        'Preparar cartas de presentación del proyecto'
    ]
    for i, acc in enumerate(acciones, 1):
        doc.add_paragraph(f'{i}. {acc}')
    
    # Archivos
    doc.add_heading('8. Archivos Disponibles', level=1)
    doc.add_paragraph('Shapefiles para ArcMap:')
    doc.add_paragraph('  - datos/ [corredor] .shp - Bloques del corredor')
    doc.add_paragraph('  - datos/ localidades_osm.shp - Localidades OSM')
    doc.add_paragraph('  - analisis/ ejidos_cercanos.csv - Lista de ejidos')
    doc.add_paragraph('  - analisis/ localidades_cercanas.csv - Lista de localidades')
    
    doc.save(output_path)
    print(f'Guardado: {output_path}')

# ============== MAIN ==============

print("Generando reportes actualizados...")

# Tampico-San Luis
crear_reporte(
    'Tampico - San Luis Potosí',
    os.path.join(BASE_DIR, "Tampico_SanLuis", "datos", "Tampico_SanLuis-50km_blocks.shp"),
    os.path.join(BASE_DIR, "Tampico_SanLuis", "analisis", "ejidos_Tampico_SanLuis.csv"),
    os.path.join(BASE_DIR, "Tampico_SanLuis", "analisis", "localidades_cercanas.csv"),
    os.path.join(BASE_DIR, "Tampico_SanLuis", "ANALISIS_DEL_PROYECTO.docx"),
    'Corredor de canalización de ~300+ km en Tamaulipas y San Luis Potosí.',
    {
        'Tamaulipas': {'cve': '30', 'municipios': 'Aldama, Güémez, San Fernando, Valle Hermoso'},
        'San Luis Potosí': {'cve': '24', 'municipios': 'Ciudad Valles, Tamazunchale, Ébano'}
    }
)

# Veracruz-Tampico
crear_reporte(
    'Veracruz - Tampico',
    os.path.join(BASE_DIR, "Veracruz_Tampico", "datos", "Veracruz-Tampico-50km_blocks_v4.shp"),
    os.path.join(BASE_DIR, "Veracruz_Tampico", "analisis", "ejidos_Veracruz_Tampico.csv"),
    os.path.join(BASE_DIR, "Veracruz_Tampico", "analisis", "localidades_cercanas.csv"),  # Puede no existir
    os.path.join(BASE_DIR, "Veracruz_Tampico", "ANALISIS_DEL_PROYECTO.docx"),
    'Corredor de microsoldadura y canalización de ~400+ km en Tamaulipas.',
    {
        'Tamaulipas': {'cve': '30', 'municipios': 'Tampico, Cd. Madero, San Fernando, Méndez, Burgos'}
    }
)

print("¡Reportes actualizados!")
