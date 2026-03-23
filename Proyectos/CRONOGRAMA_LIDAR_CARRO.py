"""
Generador de Cronograma y Planeacion: LiDAR Vehicular + Dron
Principal: LiDAR en carro | Secundario: Dron
"""

import geopandas as gpd
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime, timedelta

BASE_DIR = r"C:\Users\ed\.openclaw\workspace\Proyectos"

def load_blocks(path):
    """Carga shapefile y resume por bloque"""
    gdf = gpd.read_file(path)
    results = []
    for bloque in sorted(gdf['bloque'].unique()):
        block_data = gdf[gdf['bloque'] == bloque]
        results.append({
            'bloque': int(bloque),
            'segmentos': len(block_data),
            'km': block_data['km'].sum(),
        })
    return pd.DataFrame(results)

def add_color(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def create_planeacion_diaria(project_name, df, start_date, output_path):
    """Crea planeacion diaria: 1 bloque = 1 dia LiDAR + 1 dia Procesamiento"""
    
    doc = Document()
    
    title = doc.add_heading(f'Planeacion Diaria: {project_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Fecha de elaboracion: {datetime.now().strftime("%d de %B de %Y")}')
    doc.add_paragraph(f'Fecha de inicio: {start_date.strftime("%d de %B de %Y")}')
    doc.add_paragraph()
    
    doc.add_heading('Resumen', level=1)
    doc.add_paragraph(f'Bloques: {len(df)}')
    doc.add_paragraph(f'Kilometros totales: {df["km"].sum():.1f} km')
    doc.add_paragraph(f'Dias de adquisicion LiDAR: {len(df)}')
    doc.add_paragraph(f'Dias de procesamiento: {len(df)}')
    doc.add_paragraph(f'Total dias: {len(df) * 2 + 7} (incluyendo descansos)')
    doc.add_paragraph()
    
    doc.add_heading('Metodologia', level=1)
    doc.add_paragraph('PRINCIPAL: LiDAR vehicular (escaneo desde vehiculo)')
    doc.add_paragraph('SECUNDARIO: Dron (zonas donde el carro no puede acceder)')
    doc.add_paragraph()
    
    # Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Dia'
    hdr[1].text = 'Fecha'
    hdr[2].text = 'Actividad'
    hdr[3].text = 'km'
    hdr[4].text = 'Metodo'
    
    current_date = start_date
    day_counter = 1
    
    for _, row in df.iterrows():
        bloque = row['bloque']
        km = row['km']
        
        # Adquisicion LiDAR
        row_v = table.add_row()
        row_v.cells[0].text = str(day_counter)
        row_v.cells[1].text = current_date.strftime('%d/%m/%Y')
        row_v.cells[2].text = f'Bloque {bloque}: Adquisicion LiDAR'
        row_v.cells[3].text = f'{km:.1f}'
        row_v.cells[4].text = 'Carro'
        add_color(row_v.cells[4], 'D9E2F3')
        current_date += timedelta(days=1)
        day_counter += 1
        
        # Procesamiento + Dron si necesario
        row_p = table.add_row()
        row_p.cells[0].text = str(day_counter)
        row_p.cells[1].text = current_date.strftime('%d/%m/%Y')
        row_p.cells[2].text = f'Bloque {bloque}: Procesamiento + Dron backup'
        row_p.cells[3].text = '-'
        row_p.cells[4].text = 'Oficina/Dron'
        add_color(row_p.cells[4], 'E2EFDA')
        current_date += timedelta(days=1)
        day_counter += 1
        
        # Descanso cada 4 bloques
        if bloque % 4 == 0 and bloque < df['bloque'].max():
            row_d = table.add_row()
            row_d.cells[0].text = str(day_counter)
            row_d.cells[1].text = current_date.strftime('%d/%m/%Y')
            row_d.cells[2].text = 'Descanso'
            row_d.cells[3].text = '-'
            row_d.cells[4].text = '-'
            add_color(row_d.cells[2], 'FFF2CC')
            current_date += timedelta(days=1)
            day_counter += 1
    
    # Equipment
    doc.add_paragraph()
    doc.add_heading('Equipo LiDAR Vehicular', level=1)
    equip_lidar = [
        'LiDAR (Velodyne/Livox)',
        'Sistema GPS/IMU',
        'Vehiculo 4x4',
        'Laptop de adquisicion',
        'Baterias y cargadores',
        'Disco duro externo',
    ]
    for item in equip_lidar:
        doc.add_paragraph(f'• {item}')
    
    doc.add_heading('Equipo Dron (Backup)', level=1)
    equip_dron = [
        'Dron DJI M300/M350 RTK',
        'LiDAR para dron o camara 360',
        'Baterias adicionales',
        'Control remoto',
    ]
    for item in equip_dron:
        doc.add_paragraph(f'• {item}')
    
    doc.save(output_path)
    print(f'Guardado: {output_path}')
    return current_date

def create_cronograma(project_name, df, start_date, output_path):
    """Crea cronograma simple"""
    
    doc = Document()
    
    title = doc.add_heading(f'Cronograma: {project_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Inicio: {start_date.strftime("%d de %B de %Y")}')
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Dia'
    hdr[1].text = 'Fecha'
    hdr[2].text = 'Actividad'
    hdr[3].text = 'km'
    hdr[4].text = 'Metodo'
    
    current_date = start_date
    day_counter = 1
    
    for _, row in df.iterrows():
        bloque = row['bloque']
        km = row['km']
        
        # LiDAR
        row_v = table.add_row()
        row_v.cells[0].text = str(day_counter)
        row_v.cells[1].text = current_date.strftime('%d/%m/%Y')
        row_v.cells[2].text = f'Bloque {bloque}: LiDAR'
        row_v.cells[3].text = f'{km:.0f}'
        row_v.cells[4].text = 'Carro'
        add_color(row_v.cells[4], 'D9E2F3')
        current_date += timedelta(days=1)
        day_counter += 1
        
        # Proceso + Dron
        row_p = table.add_row()
        row_p.cells[0].text = str(day_counter)
        row_p.cells[1].text = current_date.strftime('%d/%m/%Y')
        row_p.cells[2].text = f'Bloque {bloque}: Proceso'
        row_p.cells[3].text = '-'
        row_p.cells[4].text = 'Oficina/Dron'
        add_color(row_p.cells[4], 'E2EFDA')
        current_date += timedelta(days=1)
        day_counter += 1
        
        # Descanso
        if bloque % 4 == 0 and bloque < df['bloque'].max():
            row_d = table.add_row()
            row_d.cells[0].text = str(day_counter)
            row_d.cells[1].text = current_date.strftime('%d/%m/%Y')
            row_d.cells[2].text = 'Descanso'
            row_d.cells[3].text = '-'
            row_d.cells[4].text = '-'
            add_color(row_d.cells[2], 'FFF2CC')
            current_date += timedelta(days=1)
            day_counter += 1
    
    doc.save(output_path)
    print(f'Guardado: {output_path}')
    return current_date

# ============== MAIN ==============

print("Generando cronograma y planeacion para LiDAR vehicular + Dron...")

ts_path = os.path.join(BASE_DIR, "Tampico_SanLuis", "datos", "Tampico_SanLuis-50km_blocks.shp")
vt_path = os.path.join(BASE_DIR, "Veracruz_Tampico", "datos", "Veracruz-Tampico-50km_blocks_v4.shp")

df_ts = load_blocks(ts_path)
df_vt = load_blocks(vt_path)

print(f"Tampico-San Luis: {len(df_ts)} bloques, {df_ts['km'].sum():.1f} km")
print(f"Veracruz-Tampico: {len(df_vt)} bloques, {df_vt['km'].sum():.1f} km")

# Project 1
start_ts = datetime(2026, 3, 30)

planeacion_ts = os.path.join(BASE_DIR, "Tampico_SanLuis", "PLANEACION_DIARIA.docx")
end_ts = create_planeacion_diaria("TAMPICO - SAN LUIS (LiDAR Vehicular)", df_ts, start_ts, planeacion_ts)

cronograma_ts = os.path.join(BASE_DIR, "Tampico_SanLuis", "CRONOGRAMA.docx")
create_cronograma("TAMPICO - SAN LUIS", df_ts, start_ts, cronograma_ts)

# Project 2
start_vt = end_ts + timedelta(days=1)

planeacion_vt = os.path.join(BASE_DIR, "Veracruz_Tampico", "PLANEACION_DIARIA.docx")
end_vt = create_planeacion_diaria("VERACRUZ - TAMPICO (LiDAR Vehicular)", df_vt, start_vt, planeacion_vt)

cronograma_vt = os.path.join(BASE_DIR, "Veracruz_Tampico", "CRONOGRAMA.docx")
create_cronograma("VERACRUZ - TAMPICO", df_vt, start_vt, cronograma_vt)

print()
print("=== RESUMEN ===")
print(f"Proyecto 1: {start_ts.strftime('%d/%m')} - {end_ts.strftime('%d/%m')}")
print(f"Proyecto 2: {start_vt.strftime('%d/%m')} - {end_vt.strftime('%d/%m')}")
print(f"Duracion total: {(end_vt - start_ts).days} dias")
