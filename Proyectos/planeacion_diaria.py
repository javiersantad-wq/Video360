"""
Generador de Planeación Diaria por Bloque para Proyectos de Vuelo
"""

import geopandas as gpd
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime, timedelta

BASE_DIR = r"C:\Users\ed\.openclaw\workspace\Proyectos"

def load_and_analyze_corridor(path):
    """Carga shapefile y calcula estadísticas por bloque"""
    gdf = gpd.read_file(path)
    
    results = []
    for idx, row in gdf.iterrows():
        name = row.get('Name', f'Segmento_{idx}')
        bloque = row.get('bloque', 1)
        km = row.get('km', 0)
        
        # Estimate flight time (minutes)
        # Assuming coverage rate of ~0.35 km²/min at 80m height, 8m/s
        # With corridor width of ~40m (buffer + swath)
        area_km2 = km * 0.04  # 40m wide corridor
        flight_time = (area_km2 / 0.3456) + 10  # +10 min setup
        
        results.append({
            'idx': idx,
            'name': name,
            'bloque': int(bloque),
            'km': float(km),
            'flight_time_min': flight_time,
            'batteries': int(flight_time / 25) + 1
        })
    
    return pd.DataFrame(results)

def create_daily_plan(df, project_name, start_date, output_path):
    """Crea documento Word con planeación diaria"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Planeacion Diaria de Vuelo: {project_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Fecha de elaboracion: {datetime.now().strftime("%d de %B de %Y")}')
    doc.add_paragraph(f'Fecha de inicio sugerida: {start_date.strftime("%d de %B de %Y")}')
    doc.add_paragraph()
    
    # Summary
    total_km = df['km'].sum()
    total_blocks = df['bloque'].nunique()
    total_time = df['flight_time_min'].sum()
    total_batteries = df['batteries'].sum()
    
    doc.add_heading('Resumen del Proyecto', level=1)
    summary_data = [
        ('Total de bloques de trabajo', f'{total_blocks}'),
        ('Total de kilometros', f'{total_km:.1f} km'),
        ('Tiempo total de vuelo estimado', f'{total_time/60:.1f} horas'),
        ('Baterias necesarias (estimado)', f'{total_batteries}'),
        ('Dias de vuelo estimados (4 bloques/dia)', f'{total_blocks/4:.1f} dias'),
    ]
    for label, value in summary_data:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(value)
    
    doc.add_page_break()
    
    # Daily plan by bloque
    doc.add_heading('Dias de Trabajo por Bloque', level=1)
    
    # Group by bloque
    bloques = sorted(df['bloque'].unique())
    
    current_date = start_date
    day_counter = 1
    
    # Capacity per day (in terms of blocks)
    blocks_per_day = 4
    
    for i, bloque_num in enumerate(bloques):
        bloque_data = df[df['bloque'] == bloque_num]
        
        # Check if we need a new day
        blocks_in_current_day = (i % blocks_per_day)
        if blocks_in_current_day == 0 and i > 0:
            # Add rest day every 6 working days
            if day_counter % 6 == 5:
                # Rest day
                rest_date = current_date
                current_date += timedelta(days=1)
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run(f'--- DIA DE DESCANSO ---').bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f'Date: {rest_date.strftime("%A, %d de %B de %Y")}')
                doc.add_paragraph('Recuperacion, mantenimiento de equipo, respaldo de datos')
                doc.add_paragraph()
                current_date += timedelta(days=1)
            
            current_date += timedelta(days=1)
            day_counter += 1
        
        bloque_km = bloque_data['km'].sum()
        bloque_time = bloque_data['flight_time_min'].sum()
        bloque_batteries = bloque_data['batteries'].sum()
        
        # Day header
        doc.add_heading(f'Bloque {bloque_num} - {current_date.strftime("%d %b")}', level=2)
        
        # Create table for bloque
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Segmento'
        hdr[1].text = 'Nombre'
        hdr[2].text = 'km'
        hdr[3].text = 'Tiempo min'
        
        for _, row in bloque_data.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['idx'] + 1)
            cells[1].text = str(row['name'])[:40]
            cells[2].text = f"{row['km']:.1f}"
            cells[3].text = f"{row['flight_time_min']:.0f}"
        
        # Bloque summary
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'Total bloque {bloque_num}: ').bold = True
        p.add_run(f'{bloque_km:.1f} km, {bloque_time:.0f} min, ~{bloque_batteries} baterias')
        
        doc.add_paragraph()
    
    # Final summary table
    doc.add_page_break()
    doc.add_heading('Calendario Resumido', level=1)
    
    cal_table = doc.add_table(rows=1, cols=5)
    cal_table.style = 'Table Grid'
    hdr = cal_table.rows[0].cells
    hdr[0].text = 'Dia'
    hdr[1].text = 'Fecha'
    hdr[2].text = 'Bloque'
    hdr[3].text = 'km'
    hdr[4].text = 'Notas'
    
    current_date = start_date
    day_num = 1
    
    for i, bloque_num in enumerate(bloques):
        bloque_data = df[df['bloque'] == bloque_num]
        bloque_km = bloque_data['km'].sum()
        
        # Add rest day every 6 working days
        if day_num > 1 and (day_num - 1) % 6 == 0:
            current_date += timedelta(days=1)
            cells = cal_table.add_row().cells
            cells[0].text = str(day_num)
            cells[1].text = current_date.strftime('%d/%m/%Y')
            cells[2].text = 'DESCANSO'
            cells[3].text = '-'
            cells[4].text = 'Mantenimiento'
            day_num += 1
            current_date += timedelta(days=1)
        
        cells = cal_table.add_row().cells
        cells[0].text = str(day_num)
        cells[1].text = current_date.strftime('%d/%m/%Y')
        cells[2].text = f'Bloque {bloque_num}'
        cells[3].text = f'{bloque_km:.1f}'
        cells[4].text = f'{len(bloque_data)} segmentos'
        
        day_num += 1
        if day_num % blocks_per_day == 0 or i == len(bloques) - 1:
            current_date += timedelta(days=1)
    
    # Equipment needed
    doc.add_page_break()
    doc.add_heading('Equipo Necesario por Dia', level=1)
    
    equip_text = [
        '2 Drones con LiDAR (1 principal + 1 backup)',
        '8 Baterias por dron (4 cargando, 4 volando)',
        '2 RTK (1 base + 1 rover)',
        '2 Tablets con flight plan cargado',
        '2 Laptops para procesamiento',
        '2 Discos duros externos (backup)',
        'SD Cards: 4 por dron',
        'Kit de herramientas basico',
        'Botiquin de primeros aux',
        'Radio comunicacion (2 unidades)',
        'Agua y snacks (zona rural)',
        'Cargador de 12V para vehiculo',
    ]
    
    for item in equip_text:
        doc.add_paragraph(f'• {item}')
    
    # Notes
    doc.add_heading('Notas Importantes', level=1)
    notes = [
        'Verificar weather forecast cada manana antes de volar',
        'Respaldar datos al menos 2 veces al dia',
        'Verificar linea de vista RTK antes de cada vuelo',
        'Documentar cualquier incidencia en bitacora',
        'Coordinar con autoridades locales 1 dia antes',
        'Tiempo estimado incluye setup y breakdown',
    ]
    for note in notes:
        doc.add_paragraph(f'• {note}')
    
    doc.save(output_path)
    print(f'Guardado: {output_path}')
    return doc

# ============== MAIN ==============

print("Generando planeacion diaria...")

# Project 1: Tampico - San Luis
corredor_ts = os.path.join(BASE_DIR, "Tampico_SanLuis", "datos", "Tampico_SanLuis-50km_blocks.shp")
df_ts = load_and_analyze_corridor(corredor_ts)

# Start date: April 1, 2026
start_date_ts = datetime(2026, 4, 1)

output_ts = os.path.join(BASE_DIR, "Tampico_SanLuis", "PLANEACION_DIARIA.docx")
create_daily_plan(df_ts, "Tampico - San Luis Potosi", start_date_ts, output_ts)

# Project 2: Veracruz - Tampico
corredor_vt = os.path.join(BASE_DIR, "Veracruz_Tampico", "datos", "Veracruz-Tampico-50km_blocks_v4.shp")
df_vt = load_and_analyze_corridor(corredor_vt)

# Start date: April 5, 2026 (after finishing first project setup)
start_date_vt = datetime(2026, 4, 7)

output_vt = os.path.join(BASE_DIR, "Veracruz_Tampico", "PLANEACION_DIARIA.docx")
create_daily_plan(df_vt, "Veracruz - Tampico", start_date_vt, output_vt)

print()
print("=== RESUMEN ===")
print(f"Tampico-San Luis: {len(df_ts)} segmentos, {df_ts['km'].sum():.1f} km, {df_ts['bloque'].nunique()} bloques")
print(f"Veracruz-Tampico: {len(df_vt)} segmentos, {df_vt['km'].sum():.1f} km, {df_vt['bloque'].nunique()} bloques")
