"""
Genera Calendario Maestro de Ambos Proyectos
"""

import geopandas as gpd
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime, timedelta

BASE_DIR = r"C:\Users\ed\.openclaw\workspace\Proyectos"

def load_and_summarize(path):
    """Carga shapefile y resume por bloque"""
    gdf = gpd.read_file(path)
    results = []
    for bloque in sorted(gdf['bloque'].unique()):
        block_data = gdf[gdf['bloque'] == bloque]
        results.append({
            'bloque': int(bloque),
            'segmentos': len(block_data),
            'km': block_data['km'].sum(),
        })
    return pd.DataFrame(results)

# Load both projects
ts_path = os.path.join(BASE_DIR, "Tampico_SanLuis", "datos", "Tampico_SanLuis-50km_blocks.shp")
vt_path = os.path.join(BASE_DIR, "Veracruz_Tampico", "datos", "Veracruz-Tampico-50km_blocks_v4.shp")

df_ts = load_and_summarize(ts_path)
df_vt = load_and_summarize(vt_path)

# Create master calendar document
doc = Document()

# Title
title = doc.add_heading('Calendario Maestro: Proyectos de Vuelo', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'Fecha de elaboracion: {datetime.now().strftime("%d de %B de %Y")}')
doc.add_paragraph()

# Overall summary
doc.add_heading('Resumen General', level=1)

total_km = df_ts['km'].sum() + df_vt['km'].sum()
total_blocks_ts = len(df_ts)
total_blocks_vt = len(df_vt)

p = doc.add_paragraph()
p.add_run('Duracion total estimada: ').bold = True
p.add_run(f'{total_blocks_ts + total_blocks_vt + 4} dias (incluyendo descansos y setup)')

p = doc.add_paragraph()
p.add_run('Kilometros totales: ').bold = True
p.add_run(f'{total_km:.1f} km')

# Calendar table
doc.add_page_break()
doc.add_heading('Calendario Detallado por Proyecto', level=1)

# Project 1
doc.add_heading('PROYECTO 1: Tampico - San Luis Potosi', level=2)
doc.add_paragraph(f'Fecha de inicio: 1 de Abril de 2026')
doc.add_paragraph(f'Bloques: {total_blocks_ts}')
doc.add_paragraph(f'Kilometros: {df_ts["km"].sum():.1f} km')

# Table for Project 1
table1 = doc.add_table(rows=1, cols=5)
table1.style = 'Table Grid'
hdr = table1.rows[0].cells
hdr[0].text = 'Dia'
hdr[1].text = 'Fecha'
hdr[2].text = 'Bloque'
hdr[3].text = 'km'
hdr[4].text = 'Segmentos'

current_date = datetime(2026, 4, 1)
day_num = 1
blocks_done = 0

for _, row in df_ts.iterrows():
    # Check if rest day
    if day_num > 1 and (day_num - 1) % 6 == 0:
        cells = table1.add_row().cells
        cells[0].text = str(day_num)
        cells[1].text = current_date.strftime('%d/%m/%Y')
        cells[2].text = 'DESCANSO'
        cells[3].text = '-'
        cells[4].text = '-'
        day_num += 1
        current_date += timedelta(days=1)
    
    cells = table1.add_row().cells
    cells[0].text = str(day_num)
    cells[1].text = current_date.strftime('%d/%m/%Y')
    cells[2].text = f'Bloque {row["bloque"]}'
    cells[3].text = f'{row["km"]:.1f}'
    cells[4].text = str(row['segmentos'])
    
    day_num += 1
    blocks_done += 1
    current_date += timedelta(days=1)

# Gap between projects
gap_date = current_date
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run(f'Transicion y setup entre proyectos: {gap_date.strftime("%d de %B")} (1 dia)').bold = True
current_date += timedelta(days=1)

# Project 2
doc.add_paragraph()
doc.add_heading('PROYECTO 2: Veracruz - Tampico', level=2)
doc.add_paragraph(f'Fecha de inicio: {current_date.strftime("%d de %B de %Y")}')
doc.add_paragraph(f'Bloques: {total_blocks_vt}')
doc.add_paragraph(f'Kilometros: {df_vt["km"].sum():.1f} km')

# Table for Project 2
table2 = doc.add_table(rows=1, cols=5)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Dia'
hdr2[1].text = 'Fecha'
hdr2[2].text = 'Bloque'
hdr2[3].text = 'km'
hdr2[4].text = 'Segmentos'

day_num = 1

for _, row in df_vt.iterrows():
    # Check if rest day
    if day_num > 1 and (day_num - 1) % 6 == 0:
        cells = table2.add_row().cells
        cells[0].text = str(day_num + 20)  # Continue from previous
        cells[1].text = current_date.strftime('%d/%m/%Y')
        cells[2].text = 'DESCANSO'
        cells[3].text = '-'
        cells[4].text = '-'
        day_num += 1
        current_date += timedelta(days=1)
    
    cells = table2.add_row().cells
    cells[0].text = str(day_num + 20)
    cells[1].text = current_date.strftime('%d/%m/%Y')
    cells[2].text = f'Bloque {row["bloque"]}'
    cells[3].text = f'{row["km"]:.1f}'
    cells[4].text = str(row['segmentos'])
    
    day_num += 1
    current_date += timedelta(days=1)

# Final date
final_date = current_date - timedelta(days=1)

# Phase summary
doc.add_paragraph()
doc.add_heading('Fases del Proyecto', level=1)

phases = [
    ('Fase 1: Preparacion', '1-2 dias', 'Oficina: revision de equipos, descarga de datos, planning'),
    ('Fase 2: GCPs y Campo', '2-3 dias', 'Medicion de puntos de control terrestre'),
    ('Fase 3: Vuelo Tampico-SLP', f'{20} dias', f'{total_blocks_ts} bloques, {df_ts["km"].sum():.1f} km'),
    ('Fase 4: Transicion', '1 dia', 'Movimiento de equipo, setup en nueva zona'),
    ('Fase 5: Vuelo Veracruz-Tamps', f'{total_blocks_vt + 2} dias', f'{total_blocks_vt} bloques, {df_vt["km"].sum():.1f} km'),
    ('Fase 6: Procesamiento', '5-7 dias', 'Procesamiento de datos, QA, entrega'),
]

for phase_name, duration, desc in phases:
    p = doc.add_paragraph()
    p.add_run(f'{phase_name} ({duration}): ').bold = True
    p.add_run(desc)

# Equipment list
doc.add_page_break()
doc.add_heading('Equipo Requerido', level=1)

doc.add_heading('Drones y Sensores', level=2)
equipo_drones = [
    'DJI M350 RTK (o equivalente) - 2 unidades',
    'LiDAR sensor (Velodyne/Livox/DJI L1) - 1 unidad principal',
    'Camara 360 (Insta360 X4) - 1 unidad',
    'Dron backup (Mavic 3) - 1 unidad',
]
for item in equipo_drones:
    doc.add_paragraph(f'• {item}')

doc.add_heading('Georreferenciacion', level=2)
equipo_geo = [
    'RTK Base - CHC Nav i93 o equivalente',
    'RTK Rover - CHC Nav i93 o equivalente',
    'GPS de alta precision para GCPs',
]
for item in equipo_geo:
    doc.add_paragraph(f'• {item}')

doc.add_heading('Apoyo', level=2)
equipo_apoyo = [
    'Laptop procesamiento - i7+, 32GB RAM, 1TB SSD',
    'Disco duro externo - 2TB+ (2 unidades)',
    'Tablet flight planning - iPad o similar',
    'Radios comunicacion - 2 unidades',
    'Vehiculo de apoyo - 4x4 recomendado',
    'Kit primeros aux',
]
for item in equipo_apoyo:
    doc.add_paragraph(f'• {item}')

# Save
output = os.path.join(BASE_DIR, "CALENDARIO_MAESTRO.docx")
doc.save(output)
print(f'Guardado: {output}')

# Also create a summary CSV
summary_data = []
current_date = datetime(2026, 4, 1)

for _, row in df_ts.iterrows():
    if day_num > 1 and (day_num - 1) % 6 == 0:
        summary_data.append({
            'Dia': day_num,
            'Fecha': current_date.strftime('%Y-%m-%d'),
            'Proyecto': 'DESCANSO',
            'Bloque': '-',
            'km': 0,
            'Segmentos': 0
        })
        day_num += 1
        current_date += timedelta(days=1)
    
    summary_data.append({
        'Dia': day_num,
        'Fecha': current_date.strftime('%Y-%m-%d'),
        'Proyecto': 'Tampico-San Luis',
        'Bloque': row['bloque'],
        'km': row['km'],
        'Segmentos': row['segmentos']
    })
    day_num += 1
    current_date += timedelta(days=1)

# Gap day
summary_data.append({
    'Dia': day_num,
    'Fecha': current_date.strftime('%Y-%m-%d'),
    'Proyecto': 'TRANSICIÓN',
    'Bloque': '-',
    'km': 0,
    'Segmentos': 0
})
day_num += 1
current_date += timedelta(days=1)

for _, row in df_vt.iterrows():
    if day_num > 1 and (day_num - 1) % 6 == 0:
        summary_data.append({
            'Dia': day_num,
            'Fecha': current_date.strftime('%Y-%m-%d'),
            'Proyecto': 'DESCANSO',
            'Bloque': '-',
            'km': 0,
            'Segmentos': 0
        })
        day_num += 1
        current_date += timedelta(days=1)
    
    summary_data.append({
        'Dia': day_num,
        'Fecha': current_date.strftime('%Y-%m-%d'),
        'Proyecto': 'Veracruz-Tampico',
        'Bloque': row['bloque'],
        'km': row['km'],
        'Segmentos': row['segmentos']
    })
    day_num += 1
    current_date += timedelta(days=1)

summary_df = pd.DataFrame(summary_data)
csv_output = os.path.join(BASE_DIR, "CALENDARIO_RESUMEN.csv")
summary_df.to_csv(csv_output, index=False)
print(f'CSV guardado: {csv_output}')

print()
print("Dias totales:", len(summary_df))
print("Dias de vuelo:", len(summary_df[summary_df['Proyecto'] != 'DESCANSO']))
print("Fecha final estimada:", (current_date - timedelta(days=1)).strftime('%d de %B de %Y'))
