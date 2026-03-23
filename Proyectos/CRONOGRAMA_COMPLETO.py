"""
Generador de Planeacion Diaria y Cronograma - Formato correcto
1 bloque = 1 dia de VUELO + 1 dia de PROCESO
"""

import geopandas as gpd
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime, timedelta

BASE_DIR = r"C:\Users\ed\.openclaw\workspace\Proyectos"

def load_blocks(path):
    """Carga shapefile y resume por bloque"""
    gdf = gpd.read_file(path)
    results = []
    for bloque in sorted(gdf['bloque'].unique()):
        block_data = gdf[gdf['bloque'] == bloque]
        results.append({
            'bloque': int(bloque),
            'segmentos': len(block_data),
            'km': block_data['km'].sum(),
            'names': [str(n) for n in block_data['Name'].dropna().tolist()]
        })
    return pd.DataFrame(results)

def add_color(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def create_planeacion_diaria(project_name, df, start_date, output_path):
    """Crea documento de planeacion diaria - 1 bloque por dia"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Planeacion Diaria: {project_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Fecha de elaboracion: {datetime.now().strftime("%d de %B de %Y")}')
    doc.add_paragraph(f'Fecha de inicio: {start_date.strftime("%d de %B de %Y")}')
    doc.add_paragraph()
    
    # Summary
    total_km = df['km'].sum()
    total_bloques = len(df)
    
    doc.add_heading('Resumen', level=1)
    doc.add_paragraph(f'Bloques: {total_bloques}')
    doc.add_paragraph(f'Kilometros totales: {total_km:.1f} km')
    doc.add_paragraph(f'Dias de vuelo: {total_bloques}')
    doc.add_paragraph(f'Dias de procesamiento: {total_bloques}')
    doc.add_paragraph(f'Total dias: {total_bloques * 2}')
    
    doc.add_paragraph()
    
    # Daily detail
    current_date = start_date
    day_counter = 1
    
    doc.add_heading('Detalle por Dia', level=1)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Dia'
    hdr[1].text = 'Fecha'
    hdr[2].text = 'Actividad'
    hdr[3].text = 'km'
    hdr[4].text = 'Segmentos'
    
    for _, row in df.iterrows():
        bloque = row['bloque']
        km = row['km']
        segs = row['segmentos']
        
        # Vuelo
        row_v = table.add_row()
        row_v.cells[0].text = str(day_counter)
        row_v.cells[1].text = current_date.strftime('%d/%m/%Y')
        row_v.cells[2].text = f'Bloque {bloque}: Vuelo'
        row_v.cells[3].text = f'{km:.1f}'
        row_v.cells[4].text = str(segs)
        add_color(row_v.cells[2], 'D9E2F3')
        current_date += timedelta(days=1)
        day_counter += 1
        
        # Proceso
        row_p = table.add_row()
        row_p.cells[0].text = str(day_counter)
        row_p.cells[1].text = current_date.strftime('%d/%m/%Y')
        row_p.cells[2].text = f'Bloque {bloque}: Procesamiento'
        row_p.cells[3].text = '-'
        row_p.cells[4].text = '-'
        add_color(row_p.cells[2], 'E2EFDA')
        current_date += timedelta(days=1)
        day_counter += 1
    
    # Equipment
    doc.add_paragraph()
    doc.add_heading('Equipo por Dia', level=1)
    equip = [
        '2 Drones (1 principal LiDAR + 1 backup)',
        '8 Baterias por dron',
        '2 RTK (base + rover)',
        '2 Laptops',
        '2 Discos duros externos',
        'SD Cards',
    ]
    for item in equip:
        doc.add_paragraph(f'• {item}')
    
    doc.save(output_path)
    print(f'Guardado: {output_path}')
    return current_date

def create_cronograma_gantt(df, project_name, start_date, output_path):
    """Crea cronograma tipo Gantt"""
    
    doc = Document()
    
    title = doc.add_heading(f'Cronograma: {project_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Inicio: {start_date.strftime("%d de %B de %Y")}')
    doc.add_paragraph()
    
    # Table with dates as columns
    # Calculate number of days needed
    num_days = len(df) * 2 + 2  # 2 days per bloque + buffer
    
    current_date = start_date
    dates = []
    for i in range(num_days):
        dates.append(current_date)
        current_date += timedelta(days=1)
    
    # Create table: rows = actividades, cols = dates
    table = doc.add_table(rows=1, cols=num_days + 2)
    table.style = 'Table Grid'
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = 'Actividad'
    hdr[1].text = 'Fecha'
    for i, d in enumerate(dates):
        hdr[2 + i].text = d.strftime('%d/%m')
        for para in hdr[2 + i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(7)
    
    # Add rows for each bloque
    current_date = start_date
    for _, row in df.iterrows():
        bloque = row['bloque']
        km = row['km']
        
        # Vuelo row
        row_v = table.add_row()
        row_v.cells[0].text = f'Bloque {bloque}: Vuelo'
        row_v.cells[1].text = current_date.strftime('%d/%m')
        
        # Mark the date column
        day_idx = dates.index(current_date) if current_date in dates else -1
        if day_idx >= 0:
            for col in range(2 + day_idx, num_days + 2):
                row_v.cells[col].text = ''
                add_color(row_v.cells[col], 'D9E2F3')
        
        current_date += timedelta(days=1)
        
        # Proceso row
        row_p = table.add_row()
        row_p.cells[0].text = f'Bloque {bloque}: Proceso'
        row_p.cells[1].text = current_date.strftime('%d/%m')
        
        day_idx = dates.index(current_date) if current_date in dates else -1
        if day_idx >= 0:
            for col in range(2 + day_idx, num_days + 2):
                row_p.cells[col].text = ''
                add_color(row_p.cells[col], 'E2EFDA')
        
        current_date += timedelta(days=1)
    
    doc.save(output_path)
    print(f'Guardado: {output_path}')
    return current_date

# ============== MAIN ==============

print("Generando planeacion y cronograma corregidos...")

# Load data
ts_path = os.path.join(BASE_DIR, "Tampico_SanLuis", "datos", "Tampico_SanLuis-50km_blocks.shp")
vt_path = os.path.join(BASE_DIR, "Veracruz_Tampico", "datos", "Veracruz-Tampico-50km_blocks_v4.shp")

df_ts = load_blocks(ts_path)
df_vt = load_blocks(vt_path)

print(f"Tampico-San Luis: {len(df_ts)} bloques, {df_ts['km'].sum():.1f} km")
print(f"Veracruz-Tampico: {len(df_vt)} bloques, {df_vt['km'].sum():.1f} km")

# Start dates
start_ts = datetime(2026, 3, 30)
start_vt = None

# Project 1: Tampico - San Luis
print("\n=== PROYECTO 1: TAMPICO - SAN LUIS ===")
planeacion_ts = os.path.join(BASE_DIR, "Tampico_SanLuis", "PLANEACION_DIARIA.docx")
end_ts = create_planeacion_diaria("TAMPICO - SAN LUIS", df_ts, start_ts, planeacion_ts)

cronograma_ts = os.path.join(BASE_DIR, "Tampico_SanLuis", "CRONOGRAMA.docx")
create_cronograma_gantt(df_ts, "TAMPICO - SAN LUIS", start_ts, cronograma_ts)

# Project 2: Veracruz - Tampico
start_vt = end_ts + timedelta(days=1)
print(f"\n=== PROYECTO 2: VERACRUZ - TAMPICO ===")
print(f"Inicio proyecto 2: {start_vt.strftime('%d de %B de %Y')}")

planeacion_vt = os.path.join(BASE_DIR, "Veracruz_Tampico", "PLANEACION_DIARIA.docx")
end_vt = create_planeacion_diaria("VERACRUZ - TAMPICO", df_vt, start_vt, planeacion_vt)

cronograma_vt = os.path.join(BASE_DIR, "Veracruz_Tampico", "CRONOGRAMA.docx")
create_cronograma_gantt(df_vt, "VERACRUZ - TAMPICO", start_vt, cronograma_vt)

print("\n=== RESUMEN ===")
print(f"Proyecto 1: {start_ts.strftime('%d/%m')} - {end_ts.strftime('%d/%m')}")
print(f"Proyecto 2: {start_vt.strftime('%d/%m')} - {end_vt.strftime('%d/%m')}")
print(f"Duracion total: {(end_vt - start_ts).days} dias")
