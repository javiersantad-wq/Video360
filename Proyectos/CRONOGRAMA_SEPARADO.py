"""
Generador de Cronograma Separado - Un dia por bloque
Siguiendo el formato del documento de referencia
"""

import geopandas as gpd
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime, timedelta

BASE_DIR = r"C:\Users\ed\.openclaw\workspace\Proyectos"

def load_blocks(path):
    """Carga shapefile y resume por bloque"""
    gdf = gpd.read_file(path)
    results = []
    for bloque in sorted(gdf['bloque'].unique()):
        block_data = gdf[gdf['bloque'] == bloque]
        results.append({
            'bloque': int(bloque),
            'segmentos': len(block_data),
            'km': block_data['km'].sum(),
        })
    return pd.DataFrame(results)

def add_color_to_cell(cell, hex_color):
    """Agrega color de fondo a una celda"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_cronograma(project_name, df, start_date, output_path):
    """Crea cronograma simple tipo lista"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f'PROYECTO: {project_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Plan de vuelo - LiDAR con procesamiento')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Calculate timeline
    num_bloques = len(df)
    current_date = start_date
    
    # Create table with: Dia, Fecha, Actividad, Tipo
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Dia'
    hdr[1].text = 'Fecha'
    hdr[2].text = 'Actividad'
    hdr[3].text = 'Tipo'
    
    day_counter = 1
    
    for idx, row_data in df.iterrows():
        bloque_num = row_data['bloque']
        km = row_data['km']
        
        # Vuelo row
        row_v = table.add_row()
        row_v.cells[0].text = str(day_counter)
        row_v.cells[1].text = current_date.strftime('%d/%m/%Y')
        row_v.cells[2].text = f'Bloque {bloque_num} ({km:.0f} km): Vuelo'
        row_v.cells[3].text = 'Vuelo'
        add_color_to_cell(row_v.cells[3], 'D9E2F3')
        current_date += timedelta(days=1)
        day_counter += 1
        
        # Proceso row
        row_p = table.add_row()
        row_p.cells[0].text = str(day_counter)
        row_p.cells[1].text = current_date.strftime('%d/%m/%Y')
        row_p.cells[2].text = f'Bloque {bloque_num}: Proceso'
        row_p.cells[3].text = 'Proceso'
        add_color_to_cell(row_p.cells[3], 'E2EFDA')
        current_date += timedelta(days=1)
        day_counter += 1
    
    # Add summary at bottom
    doc.add_paragraph()
    total_days = day_counter - 1
    total_km = df['km'].sum()
    
    p = doc.add_paragraph()
    p.add_run('RESUMEN: ').bold = True
    p.add_run(f'{num_bloques} bloques, {total_days} dias de trabajo, {total_km:.1f} km')
    
    doc.save(output_path)
    print(f'Guardado: {output_path}')
    return current_date

# ============== MAIN ==============

print("Generando cronograma separado por bloque...")

# Load data
ts_path = os.path.join(BASE_DIR, "Tampico_SanLuis", "datos", "Tampico_SanLuis-50km_blocks.shp")
vt_path = os.path.join(BASE_DIR, "Veracruz_Tampico", "datos", "Veracruz-Tampico-50km_blocks_v4.shp")

df_ts = load_blocks(ts_path)
df_vt = load_blocks(vt_path)

print(f"Tampico-San Luis: {len(df_ts)} bloques, {df_ts['km'].sum():.1f} km")
print(f"Veracruz-Tampico: {len(df_vt)} bloques, {df_vt['km'].sum():.1f} km")

# Start date: March 30, 2026
start_date_ts = datetime(2026, 3, 30)

# Project 1: Tampico - San Luis
output_ts = os.path.join(BASE_DIR, "Tampico_SanLuis", "CRONOGRAMA_SEPARADO.docx")
end_date_1 = create_cronograma("TAMPICO - SAN LUIS", df_ts, start_date_ts, output_ts)

# Project 2: Veracruz - Tampico
start_date_vt = end_date_1 + timedelta(days=1)
output_vt = os.path.join(BASE_DIR, "Veracruz_Tampico", "CRONOGRAMA_SEPARADO.docx")
end_date_2 = create_cronograma("VERACRUZ - TAMPICO", df_vt, start_date_vt, output_vt)

print()
print(f"Proyecto 1: {start_date_ts.strftime('%d/%m')} - {end_date_1.strftime('%d/%m')}")
print(f"Proyecto 2: {start_date_vt.strftime('%d/%m')} - {end_date_2.strftime('%d/%m')}")
